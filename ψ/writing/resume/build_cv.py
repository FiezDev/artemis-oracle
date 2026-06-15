#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Single-source CV generator for Ittipol Vongapai.
Renders md / txt / html (-> Chrome PDF) / docx for International (full + 1-page)
and Thai-market variants. ATS-clean: single column, standard headings, selectable
text (ligatures disabled), Arial, no layout tables. Minimal layout, red name accent."""

import os, subprocess, shutil, html as _html

OUT = "/tmp/cvbuild"
DIST = os.path.join(OUT, "dist")
DEST = os.environ.get("CV_DEST", "/Users/fiez/Dev/artemis-oracle/ψ/writing/resume")
CHROME = "/Applications/Google Chrome.app/Contents/MacOS/Google Chrome"
ACCENT = "#1D4ED8"  # blue accent

# ───────────────────────────── CONTENT (single source of truth) ─────────────
NAME = "Ittipol Vongapai"
TITLE = "Senior Full-Stack Engineer — React / Next.js · Node / Bun · GraphQL · AI Systems"
CONTACT = {
    "location": "Bang Phli, Samut Prakan, Thailand",
    "phone": "+66 91 721 0274",
    "email": "itti.task@gmail.com",
    "github": "github.com/FiezDev",
    "linkedin": "linkedin.com/in/fiezdev",
    "portfolio": "fiez.dev/portfolio",
}

SUMMARY = (
    "Senior full-stack engineer with 5 years building production SaaS, multi-tenant platforms, and "
    "applied-AI systems for Thai enterprises. Currently Dev Lead of the developer team at Go Thailand, "
    "where I architect and ship an agritech IoT platform (RiceGuard) end-to-end — GraphQL API, operations dashboards, and "
    "the full AWS production infrastructure. I also build the supporting tooling — QC, dataset labeling and "
    "dashboards — for the AI team's vehicle and person-recognition models. Owner-class across the stack: "
    "PostgreSQL/Drizzle/MySQL data layers, "
    "Bun/Elysia/Node/Django APIs, and React 19/Next.js/Vue/React Native frontends. I design and build my "
    "own AI systems — LLM pipelines, multi-agent automation, and MCP servers — and the data and QC tooling "
    "that feeds the AI team's computer-vision models."
)
SUMMARY_SHORT = (
    "Senior full-stack engineer (5 yrs) building production SaaS, multi-tenant platforms, and applied-AI "
    "systems. Dev Lead of the developer team at Go Thailand — architecting the RiceGuard agritech IoT platform (GraphQL API, "
    "dashboards, full AWS infrastructure) plus the tooling behind the AI team's vehicle/person-recognition models. "
    "Owner-class across PostgreSQL/Drizzle/MySQL, Bun/Elysia/Node/Django, and React/Next.js/Vue/React Native. I build "
    "my own AI systems — LLM pipelines, multi-agent automation, MCP servers — plus the data/QC tooling for the AI team's CV models."
)

SKILLS = [
    ("Languages", "TypeScript, JavaScript, Python, PHP, C#, SQL, Dart"),
    ("Frontend", "React 19, Next.js (App Router), Vue 2/3, React Native, Flutter, Tailwind CSS, shadcn/ui, "
                 "Radix UI, Zustand, TanStack Query, React Hook Form, Zod, Vitest, Jest"),
    ("Backend", "Node.js, Bun, Elysia, Hono, Express, Django, .NET/C#, GraphQL (Apollo / Yoga), Drizzle ORM, "
                "Sequelize, REST, OAuth2, JWT"),
    ("Databases", "PostgreSQL, TimescaleDB, MySQL, MongoDB / AWS DocumentDB, Redis"),
    ("Cloud & DevOps", "AWS (EC2, RDS, S3, SES, SQS, DocumentDB), GCP, Vercel, Docker, GitHub Actions, "
                       "Infrastructure-as-Code, PM2"),
    ("AI & ML", "LLM pipelines, multi-agent automation, MCP servers, LangChain, Anthropic Claude, OpenAI, "
                "PyTorch, GroundingDINO / YOLO, pgvector / RAG"),
    ("Practices", "Multi-tenant SaaS architecture, RBAC, TDD, idempotent migrations, code review, "
                  "Agile (grooming / planning / retrospectives)"),
]
SKILLS_SHORT = [
    ("Languages", "TypeScript, JavaScript, Python, PHP, C#, SQL"),
    ("Frontend", "React 19, Next.js, Vue, React Native, Tailwind, shadcn/ui, TanStack Query, Zod"),
    ("Backend & Data", "Node.js, Bun, Elysia, Hono, Django, .NET, GraphQL, Drizzle, Sequelize · "
                       "PostgreSQL, TimescaleDB, MySQL, MongoDB, Redis"),
    ("Cloud & AI", "AWS, GCP, Vercel, Docker, IaC · LLM pipelines, multi-agent automation, MCP, "
                   "LangChain, PyTorch, pgvector"),
]

EXPERIENCE = [
    {
        "company": "Go Thailand", "role": "Dev Lead, Developer Team (Contract, Remote)",
        "dates": "January 2025 – Present",
        "lead": "Lead the developer team building production platforms for a property and CCTV-AI vertical "
                "(Go Thailand / Mobile AI Co., Ltd.), and own systems end-to-end — from data layer to "
                "dashboards to cloud infrastructure.",
        "bullets": [
            "RiceGuard — agritech IoT + AI platform: built the operations, admin and AI-Ops dashboards and "
            "the GraphQL API (8 domains on PostgreSQL 16 + TimescaleDB), and designed the full AWS production "
            "infrastructure as code — VPC across two AZs, ALB/NLB, a RabbitMQ broker for MQTTS telemetry, "
            "EC2-hosted API, S3 and Secrets Manager, with GitHub Actions CI/CD.",
            "Built the supervisor QC and dataset-labeling tooling that supports the AI team's vehicle/person-"
            "recognition models (STAR Search Engine) — human-vs-AI verification, GroundingDINO/YOLO frame "
            "extraction, and vector-similarity dedupe — feeding their PyTorch/Roboflow training.",
            "Also delivered a multi-tenant CCTV/face-search SaaS (company-scoped data isolation, GPU "
            "orchestration over H200 + H100 servers) and a property/lease/tax SaaS with Thai tax compliance "
            "and accounting-system sync.",
        ],
        "short": [
            "RiceGuard agritech IoT + AI platform: ops/admin/AI-Ops dashboards, GraphQL API "
            "(PostgreSQL + TimescaleDB), and the full AWS production infrastructure as code.",
            "Built QC + dataset tooling for the AI team's vehicle/person-recognition models (human-vs-AI QC, "
            "GroundingDINO/YOLO, vector dedupe); plus a multi-tenant CCTV/face-search SaaS and a Thai tax SaaS.",
        ],
    },
    {
        "company": "7Solutions Co., Ltd.", "role": "Senior Frontend Developer",
        "dates": "June 2024 – November 2024 · Bangkok",
        "lead": "",
        "bullets": [
            "Senior frontend on LotteryPlus (lottery web app + back office) — planned a modern React/Next.js "
            "architecture (shadcn/ui, Tailwind, React Query, Zod), built a reusable component system, and drove "
            "code quality through reviews in Agile delivery.",
        ],
        "short": [
            "Senior FE on LotteryPlus — modern React/Next.js architecture (shadcn/ui, Tailwind, React Query, "
            "Zod), reusable component system, code reviews.",
        ],
    },
    {
        "company": "NestiFly", "role": "Assistant Manager / Full-Stack Developer",
        "dates": "September 2022 – May 2024 · Bangkok",
        "lead": "",
        "bullets": [
            "Built the NestiFly CRM/API (Darkphoenix / Blackbird) end-to-end — architecture, database schema, "
            "APIs and UX/UI — and deployed the full stack to Google Cloud Platform.",
            "Maintained the Bluebird (Stocklend) React Native app across the App Store and Google Play with "
            "regular feature releases and refactors.",
        ],
        "short": [
            "Built NestiFly CRM/API (Darkphoenix/Blackbird) end-to-end, deployed to GCP; maintained the "
            "Bluebird React Native app across App Store + Google Play.",
        ],
    },
    {
        "company": "Fusion Solution Co., Ltd.", "role": "Full-Stack Developer (Contract)",
        "dates": "March 2021 – April 2022 · Bangkok",
        "lead": "",
        "bullets": [
            "Delivered client products across Vue, Flutter, AngularJS and C#/.NET — a Taxi-service app "
            "(navigation, driver ratings, admin CRUD) and a Health-Insurance web app (authentication, "
            "customer reports, coupon systems).",
        ],
        "short": [
            "Delivered client products (Taxi-service app, Health-Insurance web app) across Vue, Flutter, "
            "AngularJS and C#/.NET.",
        ],
    },
]

# Selected Projects — ONLY projects shown on the portfolio (fiez.dev/portfolio)
PROJECTS = [
    {"name": "AtEase — AI Content-Automation Platform", "meta": "Solo · production", "onepage": True,
     "bullets": ["Visual node-graph workflow builder that turns one idea into multi-platform posts — LLM "
                 "drafting, fact-check, translate, image/video generation, and auto-publishing — run in the "
                 "background by Hermes AI agents. Live behind the QOne AI News Facebook page."],
     "stack": "TypeScript · React · Bun · Node.js · LLMs · Playwright · PostgreSQL",
     "short": "Visual AI workflow builder — LLM drafting → media generation → auto-publish, run by background "
              "agents; live in production. (TypeScript · React · Bun · LLMs)"},
    {"name": "ORG-TOOLS — Internal Engineering Platform", "meta": "Team · internal", "onepage": True,
     "bullets": ["Turns Jira projects into versioned ISO documentation (LLM-drafted), with a knowledge-graph "
                 "and pgvector semantic search, an Excel/CSV issue exporter, and zero-knowledge credential "
                 "vaults (server stores only ciphertext)."],
     "stack": "TypeScript · Bun · React · PostgreSQL · Jira API · AWS S3",
     "short": "Jira → versioned ISO docs (LLM-drafted) with pgvector semantic search and zero-knowledge "
              "credential vaults. (TypeScript · Bun · React · PostgreSQL)"},
    {"name": "Vehicle Verifier — AI Classification QC Tool", "meta": "Internal", "onepage": True,
     "bullets": ["Supervisor review tool placing human labels beside AI predictions (type / color / make) "
                 "with confidence scores, auto-flagged disagreements, and a reference-image panel for visual "
                 "comparison against real vehicles."],
     "stack": "Next.js · React · TypeScript · PHP · AWS S3 · MySQL",
     "short": "Human-vs-AI vehicle-classification QC tool — confidence scores, auto-flagged disagreements, "
              "reference-image panel. (Next.js · React · TypeScript)"},
    {"name": "Image Crawler & Labelling — AI Dataset Pipeline", "meta": "Internal", "onepage": True,
     "bullets": ["End-to-end console for building AI training datasets — multi-provider crawling, MD5 + "
                 "vector-similarity dedupe, keyboard-driven review, GroundingDINO + YOLO frame extraction, "
                 "and one-click export to Roboflow; long jobs run in the background with live progress."],
     "stack": "Python · FastAPI · React · PyTorch · GroundingDINO · Roboflow · Redis · AWS S3",
     "short": "End-to-end AI dataset pipeline — crawl, vector-dedupe, label, GroundingDINO/YOLO extraction, "
              "Roboflow export. (Python · FastAPI · PyTorch)"},
]

EDUCATION = [
    "Self-taught engineer — skills built through five years of production delivery and continuous, "
    "self-directed learning across the full stack, cloud, and applied AI.",
    "Secondary education, Thailand.",
]
EDUCATION_SHORT = [
    "Self-taught engineer — five years of production delivery and continuous, self-directed learning "
    "across the full stack, cloud, and applied AI. Secondary education, Thailand.",
]

LANGUAGES = "Thai (native) · English (conversational & technical — communicates effectively)"
PERSONAL = lambda: "Date of Birth: 11 November 1983 · Nationality: Thai"

# ───────────────────────────── helpers ──────────────────────────────────────
def contact_line():
    c = CONTACT
    return f"{c['location']} · {c['phone']} · {c['email']}"

def links_line():
    c = CONTACT
    return f"GitHub: {c['github']} · LinkedIn: {c['linkedin']} · Portfolio: {c['portfolio']}"

def esc(s): return _html.escape(s)

# ───────────────────────────── Markdown ─────────────────────────────────────
def render_md(onepage=False, thai=False):
    L = [f"# {NAME}\n", f"**{TITLE}**\n", f"{contact_line()}  ", f"{links_line()}  "]
    if thai:
        L.append(f"{PERSONAL()}  ")
    L += ["", "---\n", "## Professional Summary\n", (SUMMARY_SHORT if onepage else SUMMARY) + "\n", "---\n",
          "## Core Skills\n"]
    for k, v in (SKILLS_SHORT if onepage else SKILLS):
        L.append(f"**{k}:** {v}  ")
    L += ["", "---\n", "## Professional Experience\n"]
    for e in EXPERIENCE:
        L.append(f"### {e['company']} — {e['role']}")
        L.append(f"*{e['dates']}*\n")
        if onepage:
            for b in e["short"]:
                L.append(f"- {b}")
        else:
            if e["lead"]:
                L.append(e["lead"] + "\n")
            for b in e["bullets"]:
                L.append(f"- {b}")
        L.append("")
    L += ["---\n", "## Selected Projects\n"]
    for p in PROJECTS:
        if onepage and not p["onepage"]:
            continue
        if onepage:
            L.append(f"- **{p['name']}** — {p['short']}")
        elif p.get("compact"):
            stack = f" · *{p['stack']}*" if p["stack"] else ""
            L.append(f"**{p['name']}** *({p['meta']})* — {p['bullets'][0]}{stack}\n")
        else:
            L.append(f"### {p['name']}")
            L.append(f"*{p['meta']}*\n")
            L.append(f"- {p['bullets'][0]}")
            if p["stack"]:
                L.append(f"- **Stack:** {p['stack']}")
            L.append("")
    L += ["", "---\n", "## Education & Continuous Learning\n"]
    for b in (EDUCATION_SHORT if onepage else EDUCATION):
        L.append(f"- {b}")
    L += ["", "---\n", "## Languages\n", LANGUAGES + "\n"]
    return "\n".join(L).rstrip() + "\n"

# ───────────────────────────── Plain text (ATS) ─────────────────────────────
def _wrap(s, width=98):
    out, line = [], ""
    for w in s.split():
        if len(line) + len(w) + 1 > width:
            out.append(line); line = w
        else:
            line = (line + " " + w).strip()
    if line:
        out.append(line)
    return out

def render_txt(onepage=False, thai=False):
    L = [NAME.upper(), TITLE, contact_line(), links_line()]
    if thai:
        L.append(PERSONAL())
    def hdr(t):
        L.extend(["", t.upper(), ""])
    hdr("Professional Summary"); L += _wrap(SUMMARY_SHORT if onepage else SUMMARY)
    hdr("Core Skills")
    for k, v in (SKILLS_SHORT if onepage else SKILLS):
        L += _wrap(f"{k}: {v}")
    hdr("Professional Experience")
    for e in EXPERIENCE:
        L.append(f"{e['company']} - {e['role']}"); L.append(e["dates"])
        if onepage:
            for b in e["short"]:
                L.append(f"- {b}")
        else:
            if e["lead"]:
                L += _wrap(e["lead"])
            for b in e["bullets"]:
                L.append(f"- {b}")
        L.append("")
    hdr("Selected Projects")
    for p in PROJECTS:
        if onepage and not p["onepage"]:
            continue
        if onepage:
            L.append(f"- {p['name']}: {p['short']}")
        elif p.get("compact"):
            stack = f" · Stack: {p['stack']}" if p["stack"] else ""
            L.append(f"{p['name']} ({p['meta']}) - {p['bullets'][0]}{stack}")
        else:
            L.append(f"{p['name']} ({p['meta']})")
            L.append(f"- {p['bullets'][0]}")
            if p["stack"]:
                L.append(f"- Stack: {p['stack']}")
            L.append("")
    hdr("Education & Continuous Learning")
    for b in (EDUCATION_SHORT if onepage else EDUCATION):
        L.append(f"- {b}")
    hdr("Languages"); L.append(LANGUAGES)
    return "\n".join(L).rstrip() + "\n"

# ───────────────────────────── HTML (for PDF) ───────────────────────────────
CSS = """
@page { size: A4; margin: %(margin)s; }
* { box-sizing: border-box; }
body { font-family: Arial, "Helvetica Neue", sans-serif; color:#262626;
       font-size: %(fs)spt; line-height: %(lh)s; margin:0;
       font-variant-ligatures: none;
       -webkit-font-feature-settings: "liga" 0,"clig" 0,"dlig" 0;
       font-feature-settings: "liga" 0,"clig" 0,"dlig" 0; }
.header { margin-bottom: %(hb)spx; }
.name { font-size:%(name)spt; font-weight:800; letter-spacing:1.5px; color:%(accent)s;
        text-transform:uppercase; }
.title { font-size:%(title)spt; color:#3a3a3a; font-weight:600; margin-top:3px; }
.contact { font-size:%(small)spt; color:#666; margin-top:5px; }
.photo { width:92px; height:116px; border:1.5px dashed #b3c2dd; border-radius:6px; color:#7088aa;
         display:flex; align-items:center; justify-content:center; font-size:9pt; float:right; margin:0 0 6px 12px; }
.clear { clear:both; }
h2 { font-size:%(h2)spt; text-transform:uppercase; letter-spacing:1.4px; color:#1f1f1f; font-weight:700;
     border-bottom:2px solid %(accent)s; padding-bottom:3px; margin:%(h2mt)spx 0 %(h2mb)spx; }
.exphead { display:flex; justify-content:space-between; align-items:baseline; gap:12px; margin-top:%(rolemt)spx; }
.role { font-weight:700; color:#1f1f1f; font-size:%(role)spt; }
.dates { color:#777; font-weight:500; font-size:%(small)spt; white-space:nowrap; }
.lead { margin:2px 0 3px; color:#333; }
.proj { font-weight:700; color:#1f1f1f; margin:%(rolemt)spx 0 1px; }
.pmeta { color:#888; font-style:italic; font-weight:400; font-size:%(small)spt; }
.skill b { color:#1f1f1f; }
.skill, .actline, .summary, .lang { margin:2px 0; }
ul { margin:%(ulm)spx 0 %(ulb)spx; padding-left:16px; }
li { margin:%(lim)spx 0; color:#333; }
"""

def render_html(onepage=False, thai=False):
    sizes = dict(fs=9.4, name=19, title=9.6, small=8.5, h2=9.8, role=10.6, lh=1.3,
                 margin="10mm 12mm", hb=8, h2mt=10, h2mb=4, rolemt=5, ulm=2, ulb=4, lim=1.5, accent=ACCENT) if onepage \
        else dict(fs=10, name=23, title=10.6, small=9, h2=10.4, role=11.5, lh=1.42,
                  margin="15mm 16mm", hb=12, h2mt=17, h2mb=6, rolemt=8, ulm=3, ulb=7, lim=2.5, accent=ACCENT)
    css = CSS % sizes
    P = ["<!doctype html><html><head><meta charset='utf-8'><style>%s</style></head><body>" % css, "<div class='header'>"]
    if thai:
        P.append("<div class='photo'>Photo</div>")
    P.append(f"<div class='name'>{esc(NAME)}</div>")
    P.append(f"<div class='title'>{esc(TITLE)}</div>")
    P.append(f"<div class='contact'>{esc(contact_line())}</div>")
    P.append(f"<div class='contact'>{esc(links_line())}</div>")
    if thai:
        P.append(f"<div class='contact'>{esc(PERSONAL())}</div>")
    P.append("<div class='clear'></div></div>")

    P.append("<h2>Professional Summary</h2>")
    P.append(f"<div class='summary'>{esc(SUMMARY_SHORT if onepage else SUMMARY)}</div>")

    P.append("<h2>Core Skills</h2>")
    for k, v in (SKILLS_SHORT if onepage else SKILLS):
        P.append(f"<div class='skill'><b>{esc(k)}:</b> {esc(v)}</div>")

    P.append("<h2>Professional Experience</h2>")
    for e in EXPERIENCE:
        P.append(f"<div class='exphead'><span class='role'>{esc(e['company'])} — {esc(e['role'])}</span>"
                 f"<span class='dates'>{esc(e['dates'])}</span></div>")
        if onepage:
            P.append("<ul>")
            for b in e["short"]:
                P.append(f"<li>{esc(b)}</li>")
            P.append("</ul>")
        else:
            if e["lead"]:
                P.append(f"<div class='lead'>{esc(e['lead'])}</div>")
            P.append("<ul>")
            for b in e["bullets"]:
                P.append(f"<li>{esc(b)}</li>")
            P.append("</ul>")

    P.append("<h2>Selected Projects</h2>")
    for p in PROJECTS:
        if onepage and not p["onepage"]:
            continue
        if onepage:
            P.append(f"<div class='skill'><b>{esc(p['name'])}</b> — {esc(p['short'])}</div>")
        elif p.get("compact"):
            stack = f" <span class='pmeta'>· {esc(p['stack'])}</span>" if p["stack"] else ""
            P.append(f"<div class='skill'><b>{esc(p['name'])}</b> "
                     f"<span class='pmeta'>· {esc(p['meta'])}</span> — {esc(p['bullets'][0])}{stack}</div>")
        else:
            P.append(f"<div class='proj'>{esc(p['name'])} <span class='pmeta'>· {esc(p['meta'])}</span></div><ul>")
            P.append(f"<li>{esc(p['bullets'][0])}</li>")
            if p["stack"]:
                P.append(f"<li><b>Stack:</b> {esc(p['stack'])}</li>")
            P.append("</ul>")

    P.append("<h2>Education &amp; Continuous Learning</h2><ul>")
    for b in (EDUCATION_SHORT if onepage else EDUCATION):
        P.append(f"<li>{esc(b)}</li>")
    P.append("</ul>")

    P.append("<h2>Languages</h2>")
    P.append(f"<div class='lang'>{esc(LANGUAGES)}</div>")
    P.append("</body></html>")
    return "\n".join(P)

# ───────────────────────────── DOCX (ATS) ───────────────────────────────────
def render_docx(path, onepage=False, thai=False):
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()
    for s in doc.sections:
        s.top_margin = s.bottom_margin = Inches(0.55)
        s.left_margin = s.right_margin = Inches(0.65)
    doc.styles["Normal"].font.name = "Arial"
    doc.styles["Normal"].font.size = Pt(10 if not onepage else 9.5)

    ACC = RGBColor(0x1D, 0x4E, 0xD8); INK = RGBColor(0x1f, 0x1f, 0x1f); GRY = RGBColor(0x66, 0x66, 0x66)

    def p_run(text, size, bold=False, color=None, caps=False, after=0, before=0, spacing=None):
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(after); p.paragraph_format.space_before = Pt(before)
        r = p.add_run(text); r.bold = bold; r.font.size = Pt(size)
        if color is not None:
            r.font.color.rgb = color
        if caps:
            r.font.all_caps = True
        return p

    def section(title):
        p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(3)
        r = p.add_run(title.upper()); r.bold = True; r.font.size = Pt(10.5); r.font.color.rgb = INK
        pPr = p._p.get_or_add_pPr(); pbdr = OxmlElement('w:pBdr'); b = OxmlElement('w:bottom')
        b.set(qn('w:val'), 'single'); b.set(qn('w:sz'), '12'); b.set(qn('w:space'), '2'); b.set(qn('w:color'), '1D4ED8')
        pbdr.append(b); pPr.append(pbdr)

    def body(text, bold_prefix=None, after=2):
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(after)
        if bold_prefix:
            r = p.add_run(bold_prefix); r.bold = True; r.font.color.rgb = INK
        p.add_run(text); return p

    def bullet(text, prefix=None):
        p = doc.add_paragraph(style="List Bullet"); p.paragraph_format.space_after = Pt(1.5)
        if prefix:
            r = p.add_run(prefix); r.bold = True
        p.add_run(text); return p

    p_run(NAME.upper(), 21, bold=True, color=ACC)
    p_run(TITLE, 10.5, bold=True, color=INK, after=0)
    p_run(contact_line(), 9, color=GRY, after=0)
    p_run(links_line(), 9, color=GRY, after=0)
    if thai:
        p_run(PERSONAL(), 9, color=GRY, after=0)

    section("Professional Summary"); body(SUMMARY_SHORT if onepage else SUMMARY)
    section("Core Skills")
    for k, v in (SKILLS_SHORT if onepage else SKILLS):
        body(v, bold_prefix=f"{k}: ")
    section("Professional Experience")
    for e in EXPERIENCE:
        p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(5); p.paragraph_format.space_after = Pt(0)
        r = p.add_run(f"{e['company']} — {e['role']}"); r.bold = True; r.font.color.rgb = INK
        p_run(e["dates"], 9, color=GRY, after=1)
        if onepage:
            for b in e["short"]:
                bullet(b)
        else:
            if e["lead"]:
                body(e["lead"])
            for b in e["bullets"]:
                bullet(b)
    section("Selected Projects")
    for pr in PROJECTS:
        if onepage and not pr["onepage"]:
            continue
        if onepage:
            body(pr["short"], bold_prefix=f"{pr['name']} — ")
        elif pr.get("compact"):
            stack = f"  ·  {pr['stack']}" if pr["stack"] else ""
            body(f"{pr['bullets'][0]}{stack}", bold_prefix=f"{pr['name']} — ")
        else:
            p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(0)
            r = p.add_run(pr["name"]); r.bold = True; r.font.color.rgb = INK
            r2 = p.add_run(f"  ·  {pr['meta']}"); r2.italic = True; r2.font.size = Pt(9); r2.font.color.rgb = GRY
            bullet(pr["bullets"][0])
            if pr["stack"]:
                bullet(pr["stack"], prefix="Stack: ")
    section("Education & Continuous Learning")
    for b in (EDUCATION_SHORT if onepage else EDUCATION):
        bullet(b)
    section("Languages"); body(LANGUAGES)
    doc.save(path)

# ───────────────────────────── PDF via Chrome ───────────────────────────────
def html_to_pdf(html_path, pdf_path):
    subprocess.run([CHROME, "--headless", "--disable-gpu", "--no-pdf-header-footer",
                    f"--print-to-pdf={pdf_path}", "file://" + html_path],
                   check=True, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)

# ───────────────────────────── build + deploy ───────────────────────────────
def write(path, text):
    with open(path, "w", encoding="utf-8") as f:
        f.write(text)

def main():
    os.makedirs(DIST, exist_ok=True)
    targets = [("intl-full", dict(onepage=False, thai=False)),
               ("intl-onepage", dict(onepage=True, thai=False)),
               ("th-full", dict(onepage=False, thai=True))]
    for slug, opt in targets:
        write(f"{OUT}/{slug}.md", render_md(**opt))
        write(f"{OUT}/{slug}.txt", render_txt(**opt))
        hp = f"{OUT}/{slug}.html"; write(hp, render_html(**opt))
        html_to_pdf(hp, f"{DIST}/{slug}.pdf")
    render_docx(f"{DIST}/intl-full.docx", onepage=False, thai=False)
    render_docx(f"{DIST}/intl-onepage.docx", onepage=True, thai=False)

    ddist = os.path.join(DEST, "dist"); os.makedirs(ddist, exist_ok=True)
    md_map = {"intl-full": "resume.md", "intl-onepage": "resume-onepage.md", "th-full": "resume-th.md"}
    for slug, fname in md_map.items():
        shutil.copy(f"{OUT}/{slug}.md", os.path.join(DEST, fname))
        shutil.copy(f"{OUT}/{slug}.txt", os.path.join(DEST, fname[:-3] + ".txt"))
    pdf_map = {"intl-full.pdf": "Ittipol_Vongapai_CV.pdf", "intl-onepage.pdf": "Ittipol_Vongapai_CV_1page.pdf",
               "th-full.pdf": "Ittipol_Vongapai_CV_TH.pdf", "intl-full.docx": "Ittipol_Vongapai_CV.docx",
               "intl-onepage.docx": "Ittipol_Vongapai_CV_1page.docx"}
    for src, dst in pdf_map.items():
        shutil.copy(os.path.join(DIST, src), os.path.join(ddist, dst))
    shutil.copy(os.path.join(DIST, "intl-full.pdf"), os.path.join(ddist, "ittipol_cv_v2.pdf"))

    print("BUILD OK ->", DEST)
    for f in sorted(os.listdir(ddist)):
        print("  dist/", f, os.path.getsize(os.path.join(ddist, f)), "bytes")

if __name__ == "__main__":
    main()
