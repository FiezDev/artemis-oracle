"""Generate Software Components Document (#13).

Data-driven: section 3 iterates over every frontend route discovered
(or provided via analysis["frontend_modules"]). Appendix C lists the
complete route inventory so no feature is left out.

build(config, analysis, lang="en") - lang is "en" or "th".
"""

import os
from docx import Document
import core
from strings import T


def build(config, analysis, lang="en"):
    doc = Document()
    doc_number = config.doc_number("CMP-013")
    title = T(lang, "doc13.title") + "\n" + config.PROJECT_NAME
    core.add_cover_page(doc, title, doc_number, config, lang=lang)
    core.add_revision_history(doc, config, lang=lang)

    # --- Resolve modules: prefer analysis, else use discovered routes ---
    modules = _resolve_modules(analysis, config, lang)

    # --- TOC (auto-built from modules) ---
    toc_items = [
        f"1. {T(lang, 'sec.introduction')}",
        f"   1.1 {T(lang, 'sec.purpose')}",
        f"   1.2 {T(lang, 'sec.scope')}",
        f"   1.3 {T(lang, 'sec.system_overview')}",
        f"2. {T(lang, 'sec.tech_stack')}",
        f"3. {T(lang, 'sec.frontend')}",
    ]
    for i, m in enumerate(modules, 1):
        toc_items.append(f"   3.{i} {m['title']}")
    toc_items += [
        f"4. {T(lang, 'sec.backend')}",
        f"   4.1 {T(lang, 'sec.api_services')}",
        f"   4.2 {T(lang, 'sec.db_tables')}",
        f"5. {T(lang, 'sec.infrastructure')}",
        f"6. {T(lang, 'sec.external')}",
        f"7. {T(lang, 'sec.security')}",
        T(lang, "appendix.a"),
        T(lang, "appendix.c"),
    ]
    core.add_toc(doc, items=toc_items, lang=lang)

    ss_dir = config.lang_screenshots_dir(lang) if hasattr(config, 'lang_screenshots_dir') else config.SCREENSHOTS_DIR
    if not os.path.isdir(ss_dir) or not os.listdir(ss_dir):
        ss_dir = config.SCREENSHOTS_DIR
    dg_dir = config.DIAGRAMS_DIR

    # --- 1. Introduction ---
    doc.add_heading(f"1. {T(lang, 'sec.introduction')}", level=1)
    doc.add_heading(f"1.1 {T(lang, 'sec.purpose')}", level=2)
    doc.add_paragraph(T(lang, "doc13.intro.purpose", project=config.PROJECT_NAME))
    doc.add_heading(f"1.2 {T(lang, 'sec.scope')}", level=2)
    doc.add_paragraph(T(lang, "doc13.intro.scope.heading"))
    core.add_bullet_list(doc, T(lang, "doc13.intro.scope.items"))
    doc.add_heading(f"1.3 {T(lang, 'sec.system_overview')}", level=2)
    doc.add_paragraph(_get_text(analysis, "description", lang, config.PROJECT_DESCRIPTION))
    doc.add_paragraph(
        f"{config.SYSTEM_URL or '(system URL)'} | {config.GITHUB_URL or '(GitHub URL)'}"
    )
    # High-level architecture diagram only — per-function / detail diagrams
    # live in doc 15 so no diagram is duplicated across documents.
    hl_stem = _figure(config, analysis, "system_overview", "system-overview-flow")
    hl_path = os.path.join(config.DIAGRAMS_DIR, hl_stem + ".png")
    if os.path.exists(hl_path):
        core.add_image(doc, hl_path, f"{T(lang, 'phrase.fig')} 1.1: {T(lang, 'sec.system_overview')}")
    doc.add_page_break()

    # --- 2. Technology Stack ---
    doc.add_heading(f"2. {T(lang, 'sec.tech_stack')}", level=1)
    core.add_table(doc,
        [T(lang, "phrase.layer"), T(lang, "phrase.component"),
         T(lang, "phrase.technology"), T(lang, "phrase.version_col")],
        analysis.get("tech_stack", []),
        empty_note=T(lang, "default.tech_stack"),
    )
    doc.add_page_break()

    # --- 3. Frontend Components (data-driven, one section per module) ---
    # A route belongs here only if its screenshot was captured cleanly in
    # EVERY configured language — if a page works in EN but errors in TH (or
    # vice versa), it's excluded from both language docs so section 3 stays
    # identical across languages. Skipped routes still appear in Appendix C.
    common_ids = (config.common_screenshot_ids()
                  if hasattr(config, "common_screenshot_ids") else None)
    documented = []
    for m in modules:
        if not m.get("screenshot"):
            continue
        ss_id = m["screenshot"][:-4] if m["screenshot"].endswith(".png") else m["screenshot"]
        if common_ids is not None:
            if ss_id not in common_ids:
                continue
        else:
            if not os.path.exists(os.path.join(ss_dir, m["screenshot"])):
                continue
        documented.append(m)

    doc.add_heading(f"3. {T(lang, 'sec.frontend')}", level=1)
    doc.add_paragraph(_get_text(analysis, "frontend_intro", lang,
        "This section shows every page in the system with a screenshot and description."))

    for i, m in enumerate(documented, 1):
        doc.add_heading(f"3.{i} {m['title']}", level=2)
        if m.get("route"):
            p = doc.add_paragraph()
            p.add_run(f"{T(lang, 'phrase.route')}: ").bold = True
            p.add_run(m["route"])
        if m.get("description"):
            doc.add_paragraph(m["description"])
        ss_file = m.get("screenshot")
        if ss_file:
            core.add_image(doc, os.path.join(ss_dir, ss_file),
                           f"{T(lang, 'phrase.fig')} 3.{i}: {m['title']}")
        if m.get("features"):
            doc.add_paragraph(T(lang, "phrase.key_features"))
            core.add_bullet_list(doc, m["features"])
    doc.add_page_break()

    # --- 4. Backend Components ---
    doc.add_heading(f"4. {T(lang, 'sec.backend')}", level=1)
    doc.add_heading(f"4.1 {T(lang, 'sec.api_services')}", level=2)
    api_services = analysis.get("api_services", [])
    if api_services:
        doc.add_paragraph(f"{len(api_services)} API service groups.")
    core.add_table(doc,
        [T(lang, "phrase.id"), T(lang, "phrase.service"),
         T(lang, "phrase.base_path"), T(lang, "phrase.description")],
        api_services,
        empty_note=T(lang, "default.api_services"),
    )
    doc.add_page_break()

    doc.add_heading(f"4.2 {T(lang, 'sec.db_tables')}", level=2)
    db_tables = analysis.get("db_tables", [])
    if db_tables:
        doc.add_paragraph(f"{len(db_tables)} {T(lang, 'phrase.table').lower()}s.")
    core.add_table(doc,
        [T(lang, "phrase.table"), T(lang, "phrase.description"), T(lang, "phrase.key_fields")],
        db_tables,
        empty_note=T(lang, "default.db_tables"),
    )
    doc.add_page_break()

    # --- 5. Infrastructure ---
    doc.add_heading(f"5. {T(lang, 'sec.infrastructure')}", level=1)
    core.add_table(doc,
        [T(lang, "phrase.component"), T(lang, "phrase.technology"),
         T(lang, "phrase.purpose"), T(lang, "phrase.deployment")],
        analysis.get("infrastructure", []),
        empty_note=T(lang, "default.infrastructure"),
    )
    doc.add_page_break()

    # --- 6. External Services ---
    doc.add_heading(f"6. {T(lang, 'sec.external')}", level=1)
    core.add_table(doc,
        [T(lang, "phrase.service"), T(lang, "phrase.provider"),
         T(lang, "phrase.purpose"), T(lang, "phrase.integration")],
        analysis.get("external_services", []),
        empty_note=T(lang, "default.external_services"),
    )
    doc.add_page_break()

    # --- 7. Security ---
    doc.add_heading(f"7. {T(lang, 'sec.security')}", level=1)
    core.add_table(doc,
        [T(lang, "phrase.component"), T(lang, "phrase.mechanism"), T(lang, "phrase.description")],
        analysis.get("security", []),
        empty_note=T(lang, "default.security"),
    )
    doc.add_page_break()

    # --- Appendix A: Screenshots (only those captured cleanly in every language) ---
    doc.add_heading(T(lang, "appendix.a"), level=1)
    files = []
    if os.path.isdir(ss_dir):
        files = sorted(f for f in os.listdir(ss_dir) if f.endswith(".png"))
        if common_ids is not None:
            files = [f for f in files if f[:-4] in common_ids]
    if not files:
        core.add_note(doc, T(lang, "default.appendix_a_empty"))
    else:
        for i, f in enumerate(files, 1):
            label = f.replace("-ss.png", "").replace(".png", "").replace("-", " ").title()
            doc.add_heading(f"A.{i} {label}", level=2)
            core.add_image(doc, os.path.join(ss_dir, f), f"A.{i}: {label}")

    # Appendix B (flow diagrams) intentionally omitted: detail diagrams live
    # exclusively in doc 15 to avoid duplicating the same PNG across documents.

    # --- Appendix C: Complete Route Inventory (full coverage proof) ---
    doc.add_heading(T(lang, "appendix.c"), level=1)
    doc.add_paragraph(T(lang, "doc13.appendix_c.intro"))
    route_rows = []
    for m in modules:
        route_rows.append((
            m.get("id", ""),
            m.get("title", ""),
            m.get("route", ""),
            "✓" if m.get("is_dynamic") else "",
        ))
    core.add_table(doc,
        [T(lang, "phrase.module_id"), T(lang, "phrase.name"),
         T(lang, "phrase.route"), T(lang, "phrase.dynamic_col")],
        route_rows,
        empty_note=T(lang, "default.appendix_c_empty"),
    )

    # Save
    path = config.output_path("13", lang=lang)
    doc.save(path)
    print(f"  Created: {path}")
    return path


def _figure(config, analysis, figure_key, fallback_stem):
    """Return filename stem (no extension) for a core figure with fallback."""
    figs = (getattr(config, "ARCHITECTURE", {}) or {}).get("core_figures", {})
    return figs.get(figure_key) or analysis.get(f"{figure_key}_figure") or fallback_stem


def _get_text(analysis, key, lang, default=""):
    """Read analysis[key] which may be str or {'en':..., 'th':...}."""
    val = analysis.get(key, default)
    if isinstance(val, dict):
        return val.get(lang) or val.get("en") or default
    return val or default


def _resolve_modules(analysis, config, lang):
    """Return list of dicts: {id, title, route, description, screenshot, features, is_dynamic}.

    Priority:
      1. analysis["frontend_modules"] if provided as list of dicts
      2. analysis["frontend_modules"] if provided as list of tuples (legacy 4-tuple)
      3. Discovered routes via discover.walk()
    """
    raw = analysis.get("frontend_modules", [])

    # Case 1: list of dicts
    if raw and isinstance(raw[0], dict):
        out = []
        for m in raw:
            desc = m.get(f"description_{lang}") or m.get("description_en") or m.get("description") or ""
            feats = m.get(f"features_{lang}") or m.get("features_en") or m.get("features") or []
            out.append({
                "id": m.get("id", ""),
                "title": m.get(f"title_{lang}") or m.get("title") or m.get("name", ""),
                "route": m.get("route", ""),
                "description": desc,
                "features": feats,
                "screenshot": m.get("screenshot") or _guess_screenshot(m.get("id", "")),
                "is_dynamic": m.get("is_dynamic", False),
            })
        return out

    # Case 2: list of tuples (id, name, route, description)
    if raw and isinstance(raw[0], (list, tuple)):
        out = []
        for t in raw:
            id_, name, route, desc = (list(t) + [""] * 4)[:4]
            out.append({
                "id": id_, "title": name, "route": route,
                "description": desc, "features": [],
                "screenshot": _guess_screenshot(id_),
                "is_dynamic": False,
            })
        return out

    # Case 3: discover from repo(s)
    try:
        import run as _run
        routes = _run.discover_routes(config)
        if routes:
            return [{
                "id": r["id"], "title": r.get("title_guess") or r.get("id"),
                "route": r["route"],
                "description": "", "features": [],
                "screenshot": f"{r['id']}.png",
                "is_dynamic": r.get("is_dynamic", False),
            } for r in routes]
    except Exception as e:
        print(f"  discover failed: {e}")

    return []


def _guess_screenshot(module_id):
    """Default filename convention: {id}.png — matches capture.py output."""
    return f"{module_id}.png" if module_id else ""
