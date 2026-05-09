"""Generate Software Design Document (#15).

Purpose: Complete system design with architecture diagrams, module design,
data flows, and technical specifications. Project-specific content is driven
by config.ARCHITECTURE (loaded from iso-doc.json `architecture` key) so this
document is reusable for any project.

config.ARCHITECTURE shape (all keys optional):
{
  "core_figures": {
    "system_overview":    "system-overview-flow",     # filename stem (no .png)
    "deployment":         "deployment-architecture",
    "auth_flow":          "auth-flow",
    "data_flow":          "data-flow-overview",
    "entity_relationships": "entity-relationships"
  },
  "high_level_description": "..." | {"en": "...", "th": "..."},
  "request_flow_description": "..." | {...},
  "auth_description": "..." | {...},
  "modules": [
    {
      "id": "asset-mgmt",
      "title": {"en": "Asset Management Module", "th": "..."},
      "description": {"en": "...", "th": "..."},
      "diagram": "asset-management-flow",
      "sub_sections": [
        {"heading": {...}, "description": {...},
         "table": {"headers": [...], "rows": [[...], ...]}}
      ]
    }
  ],
  "integrations": [
    {"id": "flowaccount", "title": {...}, "description": {...}, "diagram": "flowaccount-integration-flow"}
  ]
}
"""

import os
from docx import Document
from docx.shared import Pt
import core
from strings import T


def _T(node, lang, default=""):
    """Resolve a value that may be a plain string or {en, th} dict."""
    if node is None:
        return default
    if isinstance(node, dict):
        return node.get(lang) or node.get("en") or default
    return node or default


def _arch(config, analysis, key, default=None):
    """Read architecture value from config.ARCHITECTURE first, then analysis."""
    arch = getattr(config, "ARCHITECTURE", {}) or {}
    if key in arch:
        return arch[key]
    return analysis.get(key, default)


def _figure(config, analysis, figure_key, fallback_stem):
    """Return filename stem for a core figure with fallback."""
    figs = (getattr(config, "ARCHITECTURE", {}) or {}).get("core_figures", {})
    return figs.get(figure_key) or analysis.get(f"{figure_key}_figure") or fallback_stem


def build(config, analysis, lang="en"):
    doc = Document()
    doc_number = config.doc_number("DES-015")
    title = T(lang, "doc15.title") + "\n" + T(lang, "doc15.subtitle")
    core.add_cover_page(doc, title, doc_number, config, lang=lang)
    core.add_revision_history(doc, config, lang=lang)

    modules = _arch(config, analysis, "modules", []) or []
    integrations = _arch(config, analysis, "integrations", []) or []

    toc_items = [
        T(lang, "doc15.sec.1"),
        "   " + T(lang, "doc15.sec.1.1"),
        "   " + T(lang, "doc15.sec.1.2"),
        "   " + T(lang, "doc15.sec.1.3"),
        "   " + T(lang, "doc15.sec.1.4"),
        T(lang, "doc15.sec.2"),
        "   " + T(lang, "doc15.sec.2.1"),
        "   " + T(lang, "doc15.sec.2.2"),
        "   " + T(lang, "doc15.sec.2.3"),
        T(lang, "doc15.sec.3"),
        T(lang, "doc15.sec.4"),
    ]
    for i, m in enumerate(modules, 1):
        toc_items.append(f"   4.{i} {_T(m.get('title'), lang, m.get('id', ''))}")
    toc_items += [
        T(lang, "doc15.sec.5"),
        T(lang, "doc15.sec.6"),
        "   " + T(lang, "doc15.sec.6.1"),
        "   " + T(lang, "doc15.sec.6.2"),
        T(lang, "doc15.sec.7"),
        "   " + T(lang, "doc15.sec.7.1"),
        "   " + T(lang, "doc15.sec.7.2"),
        "   " + T(lang, "doc15.sec.7.3"),
        T(lang, "doc15.sec.8"),
        T(lang, "doc15.sec.9"),
        "   " + T(lang, "doc15.sec.9.1"),
        "   " + T(lang, "doc15.sec.9.2"),
        T(lang, "doc15.sec.10"),
    ]
    for i, intg in enumerate(integrations, 1):
        toc_items.append(f"   10.{i} {_T(intg.get('title'), lang, intg.get('id', ''))}")
    toc_items += [
        T(lang, "appendix.a"),
    ]
    core.add_toc(doc, lang=lang, items=toc_items)

    dg = config.DIAGRAMS_DIR
    ss = config.lang_screenshots_dir(lang) if hasattr(config, 'lang_screenshots_dir') else config.SCREENSHOTS_DIR
    if not os.path.isdir(ss) or not os.listdir(ss):
        ss = config.SCREENSHOTS_DIR

    # --- 1. Introduction ---
    doc.add_heading(T(lang, "doc15.sec.1"), level=1)
    doc.add_heading(T(lang, "doc15.sec.1.1"), level=2)
    doc.add_paragraph(T(lang, "doc15.intro.purpose", project=config.PROJECT_NAME))
    doc.add_heading(T(lang, "doc15.sec.1.2"), level=2)
    core.ensure_prose(doc,
        _T(_arch(config, analysis, "scope_description"), lang, T(lang, "doc15.intro.scope")),
        T(lang, "default.scope"))
    doc.add_heading(T(lang, "doc15.sec.1.3"), level=2)
    core.add_table(doc,
        [T(lang, "doc15.th.term"), T(lang, "doc15.th.definition")],
        analysis.get("definitions", [
            ("API", "Application Programming Interface"),
            ("JWT", "JSON Web Token - stateless authentication"),
            ("RBAC", "Role-Based Access Control"),
            ("CRUD", "Create, Read, Update, Delete"),
            ("ORM", "Object-Relational Mapping"),
            ("OTP", "One-Time Password"),
        ])
    )
    doc.add_heading(T(lang, "doc15.sec.1.4"), level=2)
    core.add_table(doc,
        [T(lang, "doc15.th.document"), T(lang, "doc15.th.number")], [
        (T(lang, "doc13.title"), config.doc_number("CMP-013")),
        (T(lang, "doc19.title"), config.doc_number("TST-019")),
        ("GitHub Repository", config.GITHUB_URL or "N/A"),
    ])
    doc.add_page_break()

    # --- 2. System Architecture ---
    doc.add_heading(T(lang, "doc15.sec.2"), level=1)

    doc.add_heading(T(lang, "doc15.sec.2.1"), level=2)
    high_level = _T(_arch(config, analysis, "high_level_description"), lang, "")
    req_flow = _T(_arch(config, analysis, "request_flow_description"), lang, "")
    if high_level or req_flow:
        if high_level:
            doc.add_paragraph(high_level)
        if req_flow:
            doc.add_paragraph(req_flow)
    else:
        core.add_note(doc, T(lang, "default.high_level"))

    doc.add_heading(T(lang, "doc15.sec.2.2"), level=2)
    core.add_image(doc, os.path.join(dg, _figure(config, analysis, "deployment", "deployment-architecture") + ".png"),
                   T(lang, "doc15.fig.2.2"))
    core.add_table(doc,
        [T(lang, "phrase.component"), T(lang, "doc15.th.container"),
         T(lang, "doc15.th.port"), T(lang, "phrase.description")],
        analysis.get("deployment", []),
        empty_note=T(lang, "default.deployment"),
    )

    doc.add_heading(T(lang, "doc15.sec.2.3"), level=2)
    core.add_table(doc,
        [T(lang, "phrase.layer"), T(lang, "phrase.component"),
         T(lang, "phrase.technology"), T(lang, "phrase.version_col")],
        analysis.get("tech_stack", []),
        empty_note=T(lang, "default.tech_stack"),
    )
    doc.add_page_break()

    # --- 3. Auth Design ---
    doc.add_heading(T(lang, "doc15.sec.3"), level=1)
    core.add_image(doc, os.path.join(dg, _figure(config, analysis, "auth_flow", "auth-flow") + ".png"),
                   T(lang, "doc15.fig.3.1"))
    core.ensure_prose(doc,
        _T(_arch(config, analysis, "auth_description"), lang, ""),
        T(lang, "default.auth"))
    doc.add_heading(T(lang, "doc15.sec.3.1"), level=2)
    core.add_table(doc,
        [T(lang, "phrase.role"), T(lang, "phrase.permissions"), T(lang, "phrase.access_level")],
        analysis.get("roles", []),
        empty_note=T(lang, "default.roles"),
    )
    doc.add_page_break()

    # --- 4. Module Design (driven by config.ARCHITECTURE.modules) ---
    doc.add_heading(T(lang, "doc15.sec.4"), level=1)
    doc.add_paragraph(T(lang, "doc15.module_design.intro"))

    if not modules:
        doc.add_paragraph(T(lang, "doc15.modules.empty"))

    for i, m in enumerate(modules, 1):
        m_title = _T(m.get("title"), lang, m.get("id", ""))
        doc.add_heading(f"4.{i} {m_title}", level=2)
        m_desc = _T(m.get("description"), lang, "")
        if m_desc:
            doc.add_paragraph(m_desc)
        diagram = m.get("diagram")
        if diagram:
            core.add_image(doc, os.path.join(dg, diagram + ".png"),
                           f"{T(lang, 'phrase.fig')} 4.{i}: {m_title}")
        for j, sub in enumerate(m.get("sub_sections", []) or [], 1):
            sub_heading = _T(sub.get("heading"), lang, "")
            if sub_heading:
                doc.add_heading(f"4.{i}.{j} {sub_heading}", level=3)
            sub_desc = _T(sub.get("description"), lang, "")
            if sub_desc:
                doc.add_paragraph(sub_desc)
            table = sub.get("table")
            if table and table.get("headers") and table.get("rows"):
                core.add_table(doc,
                    [_T(h, lang, h) if isinstance(h, dict) else h for h in table["headers"]],
                    [tuple(_T(cell, lang, cell) if isinstance(cell, dict) else cell
                           for cell in row)
                     for row in table["rows"]])
    doc.add_page_break()

    # --- 5. Data Flow Diagrams ---
    doc.add_heading(T(lang, "doc15.sec.5"), level=1)
    data_flow_desc = _T(_arch(config, analysis, "data_flow_description"), lang, "")
    core.ensure_prose(doc, data_flow_desc, T(lang, "default.data_flow"))
    core.add_image(doc, os.path.join(dg, _figure(config, analysis, "data_flow", "data-flow-overview") + ".png"),
                   T(lang, "doc15.fig.5.1"))
    doc.add_page_break()

    # --- 6. UI Design ---
    doc.add_heading(T(lang, "doc15.sec.6"), level=1)
    doc.add_heading(T(lang, "doc15.sec.6.1"), level=2)
    principles = _T(analysis.get("design_principles"), lang, []) or analysis.get("design_principles", [])
    if not principles:
        principles = T(lang, "default.design_principles")
    core.add_bullet_list(doc, principles)
    doc.add_heading(T(lang, "doc15.sec.6.2"), level=2)
    core.add_table(doc,
        [T(lang, "doc15.th.section"), T(lang, "phrase.route"), T(lang, "phrase.description")],
        analysis.get("navigation", []),
        empty_note=T(lang, "default.navigation"),
    )
    doc.add_page_break()

    # --- 7. Security Design ---
    doc.add_heading(T(lang, "doc15.sec.7"), level=1)
    doc.add_heading(T(lang, "doc15.sec.7.1"), level=2)
    core.ensure_prose(doc,
        _T(_arch(config, analysis, "auth_description"), lang, ""),
        T(lang, "default.auth"))
    doc.add_heading(T(lang, "doc15.sec.7.2"), level=2)
    core.add_table(doc,
        [T(lang, "phrase.role"), T(lang, "phrase.permissions"), T(lang, "phrase.access_level")],
        analysis.get("roles", []),
        empty_note=T(lang, "default.roles"),
    )
    doc.add_heading(T(lang, "doc15.sec.7.3"), level=2)
    protection = _T(analysis.get("data_protection"), lang, []) or analysis.get("data_protection", [])
    if not protection:
        protection = T(lang, "default.data_protection")
    core.add_bullet_list(doc, protection)
    doc.add_page_break()

    # --- 8. API Reference ---
    doc.add_heading(T(lang, "doc15.sec.8"), level=1)
    core.ensure_prose(doc,
        _T(_arch(config, analysis, "api_description"), lang, T(lang, "doc15.api.default_desc")),
        T(lang, "doc15.api.default_desc"))
    core.add_table(doc,
        [T(lang, "phrase.endpoint"), T(lang, "phrase.methods"),
         T(lang, "phrase.description"), T(lang, "phrase.mechanism")],
        analysis.get("api_endpoints", []),
        empty_note=T(lang, "default.api_endpoints"),
    )
    doc.add_page_break()

    # --- 9. Database Design ---
    doc.add_heading(T(lang, "doc15.sec.9"), level=1)
    doc.add_heading(T(lang, "doc15.sec.9.1"), level=2)
    core.add_image(doc, os.path.join(dg, _figure(config, analysis, "entity_relationships", "entity-relationships") + ".png"),
                   T(lang, "doc15.fig.9.1"))
    doc.add_heading(T(lang, "doc15.sec.9.2"), level=2)
    core.add_table(doc,
        [T(lang, "phrase.table"), T(lang, "doc15.th.columns"),
         T(lang, "doc15.th.relationships"), T(lang, "phrase.description")],
        analysis.get("db_schema", []),
        empty_note=T(lang, "default.db_schema"),
    )
    doc.add_page_break()

    # --- 10. Integration Design (driven by config.ARCHITECTURE.integrations) ---
    doc.add_heading(T(lang, "doc15.sec.10"), level=1)

    if not integrations:
        doc.add_paragraph(T(lang, "doc15.integrations.empty"))

    for i, intg in enumerate(integrations, 1):
        i_title = _T(intg.get("title"), lang, intg.get("id", ""))
        doc.add_heading(f"10.{i} {i_title}", level=2)
        i_desc = _T(intg.get("description"), lang, "")
        if i_desc:
            doc.add_paragraph(i_desc)
        diagram = intg.get("diagram")
        if diagram:
            core.add_image(doc, os.path.join(dg, diagram + ".png"),
                           f"{T(lang, 'phrase.fig')} 10.{i}: {i_title}")
    doc.add_page_break()

    # --- Appendix A: Screenshots (only those captured cleanly in every language) ---
    doc.add_heading(T(lang, "appendix.a"), level=1)
    doc.add_paragraph(T(lang, "doc15.appendix_a.intro"))
    common_ids = (config.common_screenshot_ids()
                  if hasattr(config, "common_screenshot_ids") else None)
    screenshots = []
    if os.path.exists(ss):
        screenshots = sorted([f for f in os.listdir(ss) if f.endswith('.png')])
        if common_ids is not None:
            screenshots = [f for f in screenshots if f[:-4] in common_ids]
    if not screenshots:
        core.add_note(doc, T(lang, "default.appendix_a_empty"))
    else:
        for i, s in enumerate(screenshots, 1):
            title = s.replace('-ss.png', '').replace('-', ' ').title()
            doc.add_heading(f"A.{i} {title}", level=2)
            core.add_image(doc, os.path.join(ss, s), f"A.{i}: {title}")

    # Appendix B (all diagrams dump) intentionally omitted: each diagram
    # already appears in its owning inline section (2.2 deployment, 3 auth,
    # 4.x modules, 5 data flow, 9.1 ER, 10.x integrations). No duplicates.

    # Save
    path = config.output_path("15", lang=lang)
    doc.save(path)
    print(f"  Created: {path}")
    return path
