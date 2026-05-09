"""Shared helper functions for all document generators."""

import os
import struct
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from strings import T


def _png_dim(path):
    """Return (width_px, height_px) of a PNG without PIL. None/None on failure."""
    try:
        with open(path, "rb") as f:
            head = f.read(24)
        if head[:8] != b"\x89PNG\r\n\x1a\n":
            return None, None
        w, h = struct.unpack(">II", head[16:24])
        return w, h
    except Exception:
        return None, None


def add_cover_page(doc, title, doc_number, config, lang="en"):
    """Add ISO-compliant cover page with 9-row metadata table."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.font.size = Pt(24)
    run.bold = True
    doc.add_paragraph()
    doc.add_paragraph()

    status = T(lang, "cover.approved") if config.STATUS == "Approved" else config.STATUS
    classification = T(lang, "cover.internal") if config.CLASSIFICATION == "Internal" else config.CLASSIFICATION

    tbl = doc.add_table(rows=9, cols=2)
    tbl.style = 'Table Grid'
    rows = [
        (T(lang, "cover.document_number"), doc_number),
        (T(lang, "cover.version"), config.VERSION),
        (T(lang, "cover.date"), config.DOC_DATE),
        (T(lang, "cover.author"), config.AUTHOR),
        (T(lang, "cover.status"), status),
        (T(lang, "cover.classification"), classification),
        (T(lang, "cover.organization"), config.ORGANIZATION),
        (T(lang, "cover.project"), config.PROJECT_NAME),
        (T(lang, "cover.description"), config.description_for(lang)),
    ]
    for i, (label, value) in enumerate(rows):
        tbl.rows[i].cells[0].text = label
        tbl.rows[i].cells[1].text = value
        r = tbl.rows[i].cells[0].paragraphs[0].runs
        if r:
            r[0].bold = True
    doc.add_page_break()


def add_revision_history(doc, config, lang="en"):
    """Add revision history table.

    Reads from config.REVISION_ENTRIES when populated (list of dicts with
    version / date / author / changes / approved_by) so external orchestrators
    (e.g. the jira-fetch iso-gen control plane) can emit a real diff per run.
    Falls back to the hardcoded v1.0 row when no entries provided.
    """
    doc.add_heading(T(lang, "rev.heading"), level=1)
    entries = getattr(config, "REVISION_ENTRIES", []) or []
    rows = entries if entries else [{
        "version": "1.0",
        "date": config.DOC_DATE,
        "author": config.AUTHOR,
        "changes": T(lang, "rev.initial"),
        "approved_by": "-",
    }]
    tbl = doc.add_table(rows=1 + len(rows), cols=5)
    tbl.style = 'Table Grid'
    headers = [
        T(lang, "rev.version"),
        T(lang, "rev.date"),
        T(lang, "rev.author"),
        T(lang, "rev.changes"),
        T(lang, "rev.approved_by"),
    ]
    for i, h in enumerate(headers):
        tbl.rows[0].cells[i].text = h
        r = tbl.rows[0].cells[i].paragraphs[0].runs
        if r:
            r[0].bold = True
    for ri, entry in enumerate(rows, start=1):
        tbl.rows[ri].cells[0].text = str(entry.get("version", ""))
        tbl.rows[ri].cells[1].text = str(entry.get("date", ""))
        tbl.rows[ri].cells[2].text = str(entry.get("author", ""))
        tbl.rows[ri].cells[3].text = str(entry.get("changes", ""))
        tbl.rows[ri].cells[4].text = str(entry.get("approved_by", "-"))
    doc.add_paragraph()


def add_toc(doc, items, lang="en"):
    """Add table of contents."""
    doc.add_heading(T(lang, "toc.heading"), level=1)
    for item in items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(2)
    doc.add_page_break()


_MAX_IMG_W_IN = 6.0
_MAX_IMG_H_IN = 7.0


def add_image(doc, filepath, caption, width=None):
    """Insert an image, sized to fit the page, with a centered italic caption
    that stays on the same page as the image.

    Sizing: Letter page with 1" margins leaves roughly 6.5" × 9" of text area.
    We cap the picture at 6.0" wide and 7.0" tall so headings, caption, and a
    bit of breathing room all fit on one page without spilling onto the next.

    Page-break safety: image paragraph is tagged `keep_with_next`, caption is
    tagged `keep_together` + a non-breaking trailing paragraph. Word will push
    both the image and its caption to the next page together if they don't fit
    where the image lands.
    """
    if not os.path.exists(filepath):
        raise FileNotFoundError(
            f"Required image missing: {filepath}\n"
            f"  Every add_image() call must resolve to a real file. "
            f"Check that capture/diagrams produced {os.path.basename(filepath)} "
            f"or fix the module/diagram ID it was looked up from."
        )

    px_w, px_h = _png_dim(filepath)
    target_w = _MAX_IMG_W_IN
    if px_w and px_h:
        ratio = px_h / px_w
        target_h = target_w * ratio
        if target_h > _MAX_IMG_H_IN:
            target_h = _MAX_IMG_H_IN
            target_w = target_h / ratio
        doc.add_picture(filepath, width=Inches(target_w))
    else:
        doc.add_picture(filepath, width=Inches(target_w if width is None else width))

    img_para = doc.paragraphs[-1]
    img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    img_para.paragraph_format.keep_with_next = True

    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.keep_together = True
    cap.paragraph_format.space_before = Pt(2)
    cap.paragraph_format.space_after = Pt(6)
    if cap.runs:
        cap.runs[0].italic = True


def add_note(doc, text, italic=False):
    """Add a fallback body paragraph under a heading so no section renders
    blank. Defaults to non-italic because the standard caller supplies
    substantive baseline prose (industry default content), not a placeholder
    apology. Pass italic=True only for genuine callouts."""
    p = doc.add_paragraph(text)
    if italic and p.runs:
        p.runs[0].italic = True
    p.paragraph_format.space_after = Pt(6)
    return p


def add_table(doc, headers, rows, empty_note=None):
    """Add formatted table with bold headers.

    If `rows` is empty and `empty_note` is provided, render the note (italic)
    instead of a header-only table so the preceding heading is never left
    with a ghost table under it.
    """
    if not rows and empty_note is not None:
        return add_note(doc, empty_note)
    tbl = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    tbl.style = 'Table Grid'
    for i, h in enumerate(headers):
        tbl.rows[0].cells[i].text = h
        r = tbl.rows[0].cells[i].paragraphs[0].runs
        if r:
            r[0].bold = True
    for i, row_data in enumerate(rows, 1):
        for j, cell in enumerate(row_data):
            tbl.rows[i].cells[j].text = str(cell)
    return tbl


def add_bullet_list(doc, items, empty_note=None):
    """Add bullet list. When `items` is empty and `empty_note` is provided,
    render the note (italic) instead so the heading above isn't orphaned."""
    if not items and empty_note is not None:
        return add_note(doc, empty_note)
    for item in items:
        doc.add_paragraph(item, style='List Bullet')


def ensure_prose(doc, text, empty_note):
    """Emit `text` as a paragraph if non-empty, else emit `empty_note` (italic).
    Guarantees the preceding heading has at least one paragraph under it."""
    if text and str(text).strip():
        doc.add_paragraph(str(text))
    else:
        add_note(doc, empty_note)
