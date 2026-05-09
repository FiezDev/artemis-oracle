"""Generate Test Report Document (#19)."""

import os
from docx import Document
from docx.shared import Pt
import core
from strings import T


def _pick_list(node, lang):
    """Resolve a list that may be a plain list or {en, th} dict of lists."""
    if node is None:
        return []
    if isinstance(node, dict):
        return node.get(lang) or node.get("en") or []
    return node


def build(config, analysis, lang="en"):
    """Generate DOC-{CODE}-TST-019."""
    doc = Document()
    doc_number = config.doc_number("TST-019")
    title = T(lang, "doc19.title") + "\n" + config.PROJECT_NAME
    core.add_cover_page(doc, title, doc_number, config, lang=lang)
    core.add_revision_history(doc, config, lang=lang)

    _toc = [
        "doc19.sec.1", "doc19.sec.1.1", "doc19.sec.1.2", "doc19.sec.1.3",
        "doc19.sec.1.4", "doc19.sec.1.5",
        "doc19.sec.2", "doc19.sec.2.1", "doc19.sec.2.2", "doc19.sec.2.3", "doc19.sec.2.4",
        "doc19.sec.3", "doc19.sec.3.1", "doc19.sec.3.2", "doc19.sec.3.3",
        "doc19.sec.4", "doc19.sec.4.1", "doc19.sec.4.2", "doc19.sec.4.3",
        "doc19.sec.4.4", "doc19.sec.4.5", "doc19.sec.4.6", "doc19.sec.4.7",
        "doc19.sec.5", "doc19.sec.5.1", "doc19.sec.5.2",
        "doc19.sec.6", "doc19.sec.6.1", "doc19.sec.6.2", "doc19.sec.6.3",
        "doc19.sec.7", "doc19.sec.7.1", "doc19.sec.7.2", "doc19.sec.7.3", "doc19.sec.7.4",
        "doc19.sec.8",
        "doc19.sec.appendix_a", "doc19.sec.appendix_c",
    ]
    toc_items = []
    for k in _toc:
        txt = T(lang, k)
        if "." in k.split("sec.")[1]:
            toc_items.append("   " + txt)
        else:
            toc_items.append(txt)
    core.add_toc(doc, lang=lang, items=toc_items)

    # --- 1. Introduction ---
    doc.add_heading(T(lang, "doc19.sec.1"), level=1)
    doc.add_heading(T(lang, "doc19.sec.1.1"), level=2)
    doc.add_paragraph(T(lang, "doc19.intro.purpose", project=config.PROJECT_NAME))
    doc.add_heading(T(lang, "doc19.sec.1.2"), level=2)
    doc.add_paragraph(T(lang, "doc19.scope.intro"))
    scope = _pick_list(analysis.get("test_scope"), lang)
    if scope:
        core.add_bullet_list(doc, scope)
    doc.add_heading(T(lang, "doc19.sec.1.3"), level=2)
    core.add_table(doc,
        [T(lang, "doc15.th.term"), T(lang, "doc15.th.definition")],
        analysis.get("definitions", [
            ("E2E", "End-to-End testing"), ("JWT", "JSON Web Token"),
            ("RBAC", "Role-Based Access Control"), ("P95", "95th percentile response time"),
            ("RPS", "Requests Per Second"), ("CRUD", "Create, Read, Update, Delete"),
        ]))
    doc.add_heading(T(lang, "doc19.sec.1.4"), level=2)
    core.add_table(doc,
        [T(lang, "doc15.th.document"), T(lang, "doc15.th.number")], [
        (T(lang, "doc13.title"), config.doc_number("CMP-013")),
        (T(lang, "doc15.title"), config.doc_number("DES-015")),
        ("GitHub Repository", config.GITHUB_URL or "N/A"),
    ])
    doc.add_heading(T(lang, "doc19.sec.1.5"), level=2)
    doc.add_paragraph(T(lang, "doc19.method.intro"))
    core.add_bullet_list(doc, T(lang, "doc19.phases"))
    doc.add_page_break()

    # --- 2. Test Environment ---
    doc.add_heading(T(lang, "doc19.sec.2"), level=1)
    doc.add_heading(T(lang, "doc19.sec.2.1"), level=2)
    core.add_table(doc,
        [T(lang, "doc19.th.component"), T(lang, "doc19.th.specification"),
         T(lang, "doc19.th.provider"), T(lang, "doc19.th.region")],
        analysis.get("test_infrastructure", []),
        empty_note=T(lang, "default.test_infra"))
    doc.add_heading(T(lang, "doc19.sec.2.2"), level=2)
    core.add_table(doc,
        [T(lang, "doc19.th.tool"), T(lang, "doc19.th.version"),
         T(lang, "doc19.th.purpose"), T(lang, "doc19.th.scope")],
        analysis.get("test_tools", []),
        empty_note=T(lang, "default.test_tools"))
    doc.add_heading(T(lang, "doc19.sec.2.3"), level=2)
    core.add_table(doc,
        [T(lang, "doc19.th.category"), T(lang, "doc19.th.count"), T(lang, "doc19.th.details")],
        analysis.get("test_data", []),
        empty_note=T(lang, "default.test_data"))
    doc.add_heading(T(lang, "doc19.sec.2.4"), level=2)
    core.add_table(doc,
        [T(lang, "doc19.th.phase"), T(lang, "doc19.th.tool"),
         T(lang, "doc19.th.status"), T(lang, "phrase.description")],
        analysis.get("test_pipeline", []),
        empty_note=T(lang, "default.test_pipeline"))
    doc.add_page_break()

    # --- 3. Test Plan Summary ---
    doc.add_heading(T(lang, "doc19.sec.3"), level=1)
    doc.add_heading(T(lang, "doc19.sec.3.1"), level=2)
    core.add_table(doc,
        [T(lang, "doc19.th.test_level"), T(lang, "doc19.th.approach"),
         T(lang, "doc19.th.automation"), T(lang, "doc19.th.target")],
        analysis.get("test_strategy", []),
        empty_note=T(lang, "default.test_strategy"))
    doc.add_heading(T(lang, "doc19.sec.3.2"), level=2)
    core.add_table(doc,
        [T(lang, "doc19.th.category"), T(lang, "doc19.th.tool"),
         T(lang, "doc19.th.total"), T(lang, "doc19.th.passed"),
         T(lang, "doc19.th.failed"), T(lang, "doc19.th.blocked"), T(lang, "doc19.th.pass_rate")],
        analysis.get("test_categories", []),
        empty_note=T(lang, "default.test_categories"))
    doc.add_heading(T(lang, "doc19.sec.3.3"), level=2)
    core.add_table(doc,
        [T(lang, "doc19.th.phase"), T(lang, "doc19.th.duration"),
         T(lang, "doc19.th.status"), T(lang, "doc19.th.cases")],
        analysis.get("test_timeline", []),
        empty_note=T(lang, "default.test_timeline"))
    doc.add_page_break()

    # --- 4. Test Results ---
    doc.add_heading(T(lang, "doc19.sec.4"), level=1)
    doc.add_heading(T(lang, "doc19.sec.4.1"), level=2)
    core.add_table(doc,
        [T(lang, "doc19.th.item"), T(lang, "doc19.th.detail")],
        analysis.get("test_summary", []),
        empty_note=T(lang, "default.test_summary"))
    doc.add_heading(T(lang, "doc19.sec.4.2"), level=2)
    core.add_table(doc,
        [T(lang, "doc19.th.module"), T(lang, "doc19.th.test_type"),
         T(lang, "doc19.th.cases"), T(lang, "doc19.th.passed"),
         T(lang, "doc19.th.failed"), T(lang, "doc19.th.notes")],
        analysis.get("frontend_results", []),
        empty_note=T(lang, "default.frontend_results"))
    doc.add_heading(T(lang, "doc19.sec.4.3"), level=2)
    core.add_table(doc,
        [T(lang, "doc19.th.module"), T(lang, "doc19.th.test_type"),
         T(lang, "doc19.th.cases"), T(lang, "doc19.th.passed"),
         T(lang, "doc19.th.failed"), T(lang, "doc19.th.notes")],
        analysis.get("backend_results", []),
        empty_note=T(lang, "default.backend_results"))
    doc.add_heading(T(lang, "doc19.sec.4.4"), level=2)
    core.add_table(doc,
        [T(lang, "doc19.th.integration_point"), T(lang, "doc19.th.method"),
         T(lang, "doc19.th.result"), T(lang, "doc19.th.details")],
        analysis.get("integration_results", []),
        empty_note=T(lang, "default.integration_results"))
    doc.add_heading(T(lang, "doc19.sec.4.5"), level=2)
    core.add_table(doc,
        [T(lang, "phrase.endpoint"), T(lang, "doc19.th.concurrency"),
         T(lang, "doc19.th.avg_ms"), T(lang, "doc19.th.p95_ms"), T(lang, "doc19.th.status")],
        analysis.get("performance_results", []),
        empty_note=T(lang, "default.performance_results"))
    doc.add_heading(T(lang, "doc19.sec.4.6"), level=2)
    core.add_table(doc,
        [T(lang, "doc19.th.test_case"), T(lang, "doc19.th.category"),
         T(lang, "doc19.th.result"), T(lang, "doc19.th.details")],
        analysis.get("security_results", []),
        empty_note=T(lang, "default.security_results"))
    doc.add_heading(T(lang, "doc19.sec.4.7"), level=2)
    core.add_table(doc,
        [T(lang, "doc19.th.route_group"), T(lang, "doc19.th.endpoints"),
         T(lang, "doc19.th.unit"), T(lang, "doc19.th.integration"),
         T(lang, "doc19.th.manual"), T(lang, "doc19.th.coverage")],
        analysis.get("api_coverage", []),
        empty_note=T(lang, "default.api_coverage"))
    doc.add_page_break()

    # --- 5. Defect Analysis ---
    doc.add_heading(T(lang, "doc19.sec.5"), level=1)
    doc.add_heading(T(lang, "doc19.sec.5.1"), level=2)
    core.add_table(doc,
        [T(lang, "doc19.th.severity"), T(lang, "doc19.th.count"),
         T(lang, "doc19.th.open"), T(lang, "doc19.th.fixed"), T(lang, "doc19.th.deferred")],
        analysis.get("defect_summary", []),
        empty_note=T(lang, "default.defect_summary"))
    doc.add_heading(T(lang, "doc19.sec.5.2"), level=2)
    core.add_table(doc,
        [T(lang, "phrase.id"), T(lang, "doc19.th.severity"),
         T(lang, "doc19.th.module"), T(lang, "phrase.description"),
         T(lang, "doc19.th.status"), T(lang, "doc19.th.resolution")],
        analysis.get("defect_details", []),
        empty_note=T(lang, "default.defect_details"))
    doc.add_page_break()

    # --- 6. Test Coverage ---
    doc.add_heading(T(lang, "doc19.sec.6"), level=1)
    doc.add_heading(T(lang, "doc19.sec.6.1"), level=2)
    core.add_table(doc,
        [T(lang, "doc19.th.component"), T(lang, "doc19.th.statements"),
         T(lang, "doc19.th.branches"), T(lang, "doc19.th.functions"), T(lang, "doc19.th.lines")],
        analysis.get("code_coverage", []),
        empty_note=T(lang, "default.code_coverage"))
    doc.add_heading(T(lang, "doc19.sec.6.2"), level=2)
    core.add_table(doc,
        [T(lang, "doc19.th.area"), T(lang, "doc19.th.cases"),
         T(lang, "doc19.th.covered"), T(lang, "doc19.th.coverage_pct")],
        analysis.get("functional_coverage", []),
        empty_note=T(lang, "default.functional_coverage"))
    doc.add_heading(T(lang, "doc19.sec.6.3"), level=2)
    core.add_table(doc,
        [T(lang, "doc19.th.module"), T(lang, "doc19.th.cases"),
         T(lang, "doc19.th.covered"), T(lang, "doc19.th.coverage_pct"),
         T(lang, "doc19.th.notes")],
        analysis.get("coverage_gaps", []),
        empty_note=T(lang, "default.coverage_gaps"))
    doc.add_page_break()

    # --- 7. Evaluation ---
    # Entry/exit/risk tables fall back to the standard baseline rows so the
    # document always carries a real, signable matrix — never an empty shell.
    doc.add_heading(T(lang, "doc19.sec.7"), level=1)
    doc.add_heading(T(lang, "doc19.sec.7.1"), level=2)
    entry_rows = analysis.get("entry_criteria") or T(lang, "default.entry_criteria_rows")
    core.add_table(doc,
        [T(lang, "doc19.th.criteria"), T(lang, "doc19.th.status"), T(lang, "doc19.th.evidence")],
        entry_rows)
    doc.add_heading(T(lang, "doc19.sec.7.2"), level=2)
    exit_rows = analysis.get("exit_criteria") or T(lang, "default.exit_criteria_rows")
    core.add_table(doc,
        [T(lang, "doc19.th.criteria"), T(lang, "doc19.th.target"),
         T(lang, "doc19.th.actual"), T(lang, "doc19.th.status")],
        exit_rows)
    doc.add_heading(T(lang, "doc19.sec.7.3"), level=2)
    risk_rows = analysis.get("risks") or T(lang, "default.risks_rows")
    core.add_table(doc,
        [T(lang, "doc19.th.risk"), T(lang, "doc19.th.likelihood"),
         T(lang, "doc19.th.impact"), T(lang, "doc19.th.mitigation")],
        risk_rows)
    doc.add_heading(T(lang, "doc19.sec.7.4"), level=2)
    p = doc.add_paragraph()
    run = p.add_run(T(lang, "doc19.go_label"))
    run.bold = True
    run.font.size = Pt(14)
    doc.add_paragraph(T(lang, "doc19.recommendation.default"))
    doc.add_page_break()

    # --- 8. Sign-Off ---
    doc.add_heading(T(lang, "doc19.sec.8"), level=1)
    core.add_table(doc,
        [T(lang, "phrase.role"), T(lang, "phrase.name"),
         T(lang, "doc19.th.signature"), T(lang, "cover.date")],
        analysis.get("signoff", [
            (T(lang, "phrase.role"), config.AUTHOR, "", config.DOC_DATE),
        ]))
    doc.add_page_break()

    # --- Appendix A: Screenshots ---
    doc.add_heading(T(lang, "doc19.sec.appendix_a"), level=1)
    ss_dir = config.lang_screenshots_dir(lang) if hasattr(config, 'lang_screenshots_dir') else config.SCREENSHOTS_DIR
    if not os.path.isdir(ss_dir) or not os.listdir(ss_dir):
        ss_dir = config.SCREENSHOTS_DIR
    common_ids = (config.common_screenshot_ids()
                  if hasattr(config, "common_screenshot_ids") else None)
    if os.path.exists(ss_dir):
        screenshots = sorted([f for f in os.listdir(ss_dir) if f.endswith('.png')])
        if common_ids is not None:
            screenshots = [f for f in screenshots if f[:-4] in common_ids]
        for i, ss in enumerate(screenshots, 1):
            doc.add_heading(f"A.{i} {ss.replace('-ss.png', '').replace('-', ' ').title()}", level=2)
            core.add_image(doc, os.path.join(ss_dir, ss), f"A.{i}: {ss}")

    # Appendix B (diagrams) intentionally omitted: doc 19 is the Test Report
    # and should carry only test-related content. Architecture and flow
    # diagrams live in doc 13 (high-level) and doc 15 (detail/per-function).

    # --- Appendix C: Detailed Test Cases ---
    doc.add_heading(T(lang, "doc19.sec.appendix_c"), level=1)
    core.add_table(doc,
        [T(lang, "doc19.th.tc_id"), T(lang, "doc19.th.test_case"),
         T(lang, "doc19.th.module"), T(lang, "doc19.th.result")],
        analysis.get("detailed_test_cases", []),
        empty_note=T(lang, "default.detailed_test_cases"))

    # Save
    path = config.output_path("19", lang=lang)
    doc.save(path)
    print(f"  Created: {path}")
    return path
